"""Word RFI document export — formal RFI pack using python-docx."""
from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


# ---------------------------------------------------------------------------
# Colour constants
# ---------------------------------------------------------------------------
COLOR_BLUE_DARK = RGBColor(0x1F, 0x38, 0x64)
COLOR_RED = RGBColor(0xC0, 0x00, 0x00)
COLOR_AMBER = RGBColor(0xC5, 0x5A, 0x11)
COLOR_GREEN = RGBColor(0x37, 0x56, 0x23)
COLOR_GREY = RGBColor(0x60, 0x60, 0x60)

PRIORITY_COLORS = {
    "P1": COLOR_RED,
    "P2": COLOR_AMBER,
    "P3": COLOR_GREEN,
}


# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------
def _add_horizontal_line(doc: Document) -> None:
    """Insert a paragraph with a bottom border acting as a horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_cell_bg(cell, hex_color: str) -> None:
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_footer(doc: Document) -> None:
    """Add a confidential footer to each page."""
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.clear()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("xBOQ.ai Tender Intelligence | Confidential")
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_GREY
    run.font.italic = True


def _priority_sort_key(rfi: dict) -> int:
    order = {"P1": 0, "P2": 1, "P3": 2}
    return order.get((rfi.get("priority") or "P3").upper(), 99)


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------
def _build_cover_page(doc: Document, payload: dict) -> None:
    """Cover page with title, project info, and horizontal line."""
    # Large centered heading
    heading = doc.add_heading("REQUEST FOR INFORMATION", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = COLOR_BLUE_DARK
        run.font.size = Pt(24)

    doc.add_paragraph()  # spacer

    project_id = payload.get("project_id", "—")
    run_date = date.today().isoformat()

    for label, value in [
        ("Project ID", project_id),
        ("Date", run_date),
        ("Prepared by", "xBOQ.ai"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = p.add_run(f"{label}: ")
        label_run.bold = True
        label_run.font.size = Pt(11)
        value_run = p.add_run(value)
        value_run.font.size = Pt(11)

    doc.add_paragraph()
    _add_horizontal_line(doc)
    doc.add_paragraph()


def _build_summary_table(doc: Document, rfis: list[dict]) -> None:
    """Summary counts table."""
    p = doc.add_heading("RFI Summary", level=2)
    for run in p.runs:
        run.font.color.rgb = COLOR_BLUE_DARK

    total = len(rfis)
    p1_count = sum(1 for r in rfis if (r.get("priority") or "").upper() == "P1")
    p2_count = sum(1 for r in rfis if (r.get("priority") or "").upper() == "P2")
    p3_count = sum(1 for r in rfis if (r.get("priority") or "").upper() == "P3")

    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"

    rows_data = [
        ("Total RFIs", str(total)),
        ("P1 (Critical)", str(p1_count)),
        ("P2 (Important)", str(p2_count)),
        ("P3 (Advisory)", str(p3_count)),
        ("Status", "Pending Response"),
    ]

    header_bg = "1F3864"
    row_bgs = ["F2F2F2", "FFFFFF", "F2F2F2", "FFFFFF", "F2F2F2"]

    for row_idx, (label, value) in enumerate(rows_data):
        row = tbl.rows[row_idx]
        # Label cell
        label_cell = row.cells[0]
        label_cell.text = label
        for para in label_cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
        _set_cell_bg(label_cell, row_bgs[row_idx])

        # Value cell
        value_cell = row.cells[1]
        value_cell.text = value
        for para in value_cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
        _set_cell_bg(value_cell, row_bgs[row_idx])

    doc.add_paragraph()


def _build_rfi_entries(doc: Document, rfis: list[dict]) -> None:
    """One section per RFI, sorted by priority."""
    sorted_rfis = sorted(rfis, key=_priority_sort_key)

    for idx, rfi in enumerate(sorted_rfis):
        priority = (rfi.get("priority") or "P3").upper()
        trade = (rfi.get("trade") or "general").upper()
        rfi_id = rfi.get("id", f"RFI-{idx + 1:03d}")
        question = rfi.get("question", "")
        evidence = rfi.get("evidence", "")

        # Section heading
        heading_text = f"RFI-{rfi_id} \u2014 {trade}"
        section_heading = doc.add_heading(heading_text, level=2)
        section_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in section_heading.runs:
            run.font.color.rgb = COLOR_BLUE_DARK
            run.font.size = Pt(12)

        # Priority badge line
        badge_para = doc.add_paragraph()
        badge_label = badge_para.add_run("Priority: ")
        badge_label.bold = True
        badge_label.font.size = Pt(10)
        badge_value = badge_para.add_run(priority)
        badge_value.bold = True
        badge_value.font.size = Pt(10)
        p_color = PRIORITY_COLORS.get(priority, COLOR_GREY)
        badge_value.font.color.rgb = p_color

        # Question
        q_para = doc.add_paragraph()
        q_label = q_para.add_run("Question: ")
        q_label.bold = True
        q_label.font.size = Pt(10)
        q_text = q_para.add_run(question)
        q_text.font.size = Pt(10)

        # Evidence
        ev_para = doc.add_paragraph()
        ev_label = ev_para.add_run("Evidence: ")
        ev_label.bold = True
        ev_label.italic = True
        ev_label.font.size = Pt(10)
        ev_text = ev_para.add_run(evidence)
        ev_text.italic = True
        ev_text.font.size = Pt(10)

        # Response field
        resp_para = doc.add_paragraph()
        resp_label = resp_para.add_run("Response: ")
        resp_label.bold = True
        resp_label.font.size = Pt(10)
        resp_blank = resp_para.add_run("___________________________________________")
        resp_blank.font.size = Pt(10)
        resp_blank.font.color.rgb = COLOR_GREY

        # Horizontal separator (not after last RFI)
        if idx < len(sorted_rfis) - 1:
            doc.add_paragraph()
            _add_horizontal_line(doc)
            doc.add_paragraph()


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------
def export_rfi_word(payload: dict, output_path: Path) -> Path:
    """Export a formal RFI Word document from the pipeline payload.

    Args:
        payload: Pipeline output dict.
        output_path: Destination .docx file path.

    Returns:
        output_path after writing.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    _add_footer(doc)

    rfis = payload.get("rfis") or []

    _build_cover_page(doc, payload)
    _build_summary_table(doc, rfis)

    if rfis:
        _build_rfi_entries(doc, rfis)
    else:
        doc.add_paragraph("No RFIs were generated for this project.")

    doc.save(str(output_path))
    return output_path
