"""
Pre-bid Meeting Agenda Builder — auto-generates a structured agenda from
review queue items and collaboration assignments.

Pure function, no Streamlit dependency. Can be tested independently.

Sprint 15: Packaging + Proof + Meeting Workflow.
"""

import io
from typing import Dict, Any, List, Optional
from datetime import datetime


# =============================================================================
# AGENDA BUILDER (pure function)
# =============================================================================

_SEVERITY_ORDER = {"high": 0, "medium": 1, "low": 2}
_TYPE_ORDER = {
    "risk_hit": 0,
    "recon_mismatch": 1,
    "conflict": 2,
    "toxic_page": 3,
    "skipped_page": 4,
}


def build_meeting_agenda(
    review_items: List[dict],
    assignments: List[dict],
    rfis: Optional[List[dict]] = None,
    assumptions: Optional[List[dict]] = None,
    project_name: str = "",
    max_items: int = 20,
) -> dict:
    """
    Build a deterministic pre-bid meeting agenda from review queue items and
    open approvals/assignments.

    Args:
        review_items: Output of build_review_queue() — sorted list of review items.
        assignments: Output of get_all_assignments() — [{entity_type, entity_id, assigned_to, due_date}].
        rfis: Optional list of RFI dicts (to identify open/draft approvals).
        assumptions: Optional list of assumption dicts.
        project_name: Project name for the agenda header.
        max_items: Maximum items per section.

    Returns:
        Deterministic agenda dict with sections and summary.
    """
    rfis = rfis or []
    assumptions = assumptions or []
    generated_at = datetime.now().isoformat()

    # Build assignment lookup: (entity_type, entity_id) -> {assigned_to, due_date}
    assignment_map = {}
    for a in assignments:
        key = (a.get("entity_type", ""), a.get("entity_id", ""))
        assignment_map[key] = {
            "assigned_to": a.get("assigned_to", ""),
            "due_date": a.get("due_date", ""),
        }

    def _enrich_item(item: dict, entity_type: str = "", entity_id: str = "") -> dict:
        """Add assignment info to an agenda item."""
        collab = assignment_map.get((entity_type, entity_id), {})
        return {
            "id": entity_id or item.get("id", item.get("source_key", "")),
            "title": item.get("title", item.get("question", "")),
            "type": item.get("type", entity_type),
            "severity": item.get("severity", "medium"),
            "assigned_to": collab.get("assigned_to", ""),
            "due_date": collab.get("due_date", ""),
            "status": item.get("status", item.get("review_status", "")),
            "recommended_action": item.get("recommended_action", "review"),
        }

    # ── Section 1: High Priority Review Items ─────────────────────────
    high_items = []
    for ri in review_items:
        if ri.get("severity") == "high":
            enriched = _enrich_item(ri, "review_item", ri.get("source_key", ""))
            high_items.append(enriched)
    high_items = high_items[:max_items]

    # ── Section 2: Open Approvals ─────────────────────────────────────
    open_approvals = []
    for rfi in rfis:
        if rfi.get("status", "draft") == "draft":
            enriched = _enrich_item(rfi, "rfi", rfi.get("id", ""))
            enriched["title"] = rfi.get("question", rfi.get("title", ""))[:100]
            enriched["status"] = "draft"
            open_approvals.append(enriched)
    for asm in assumptions:
        if asm.get("status", "") not in ("accepted", "rejected"):
            enriched = _enrich_item(asm, "assumption", asm.get("id", ""))
            enriched["title"] = asm.get("title", asm.get("assumption", ""))[:100]
            enriched["status"] = asm.get("status", "pending")
            open_approvals.append(enriched)
    from src.analysis.determinism import stable_sort_key
    open_approvals = sorted(
        open_approvals,
        key=lambda x: (x.get("type", ""), x.get("title", ""), stable_sort_key(x, ("id",))),
    )[:max_items]

    # ── Section 3: Assigned Items ─────────────────────────────────────
    assigned_items = []
    for a in assignments:
        if a.get("assigned_to"):
            assigned_items.append({
                "id": a.get("entity_id", ""),
                "title": f"[{a.get('entity_type', '').upper()}] {a.get('entity_id', '')}",
                "type": a.get("entity_type", ""),
                "severity": "",
                "assigned_to": a.get("assigned_to", ""),
                "due_date": a.get("due_date", ""),
                "status": "assigned",
                "recommended_action": "follow up",
            })
    assigned_items = sorted(
        assigned_items,
        key=lambda x: (x.get("due_date", "") or "9999", x.get("assigned_to", ""), stable_sort_key(x, ("id",))),
    )[:max_items]

    # ── Section 4: Due Soon (items with due_date set) ─────────────────
    due_soon = []
    for a in assignments:
        if a.get("due_date"):
            due_soon.append({
                "id": a.get("entity_id", ""),
                "title": f"[{a.get('entity_type', '').upper()}] {a.get('entity_id', '')}",
                "type": a.get("entity_type", ""),
                "severity": "",
                "assigned_to": a.get("assigned_to", ""),
                "due_date": a.get("due_date", ""),
                "status": "due",
                "recommended_action": "complete by due date",
            })
    due_soon = sorted(due_soon, key=lambda x: (x.get("due_date", ""), stable_sort_key(x, ("id",))))[:max_items]

    # ── Assemble sections ─────────────────────────────────────────────
    sections = []
    if high_items:
        sections.append({"title": "High Priority Review Items", "items": high_items})
    if open_approvals:
        sections.append({"title": "Open Approvals", "items": open_approvals})
    if assigned_items:
        sections.append({"title": "Assigned Items", "items": assigned_items})
    if due_soon:
        sections.append({"title": "Due Soon", "items": due_soon})

    # If no items at all, add an empty section for schema consistency
    if not sections:
        sections.append({"title": "No Agenda Items", "items": []})

    # Summary
    all_items = high_items + open_approvals + assigned_items + due_soon
    high_count = sum(1 for i in all_items if i.get("severity") == "high")
    assigned_count = sum(1 for i in all_items if i.get("assigned_to"))

    return {
        "project_name": project_name,
        "generated_at": generated_at,
        "sections": sections,
        "summary": {
            "total_items": len(all_items),
            "high_count": high_count,
            "assigned_count": assigned_count,
            "sections_count": len(sections),
        },
    }


# =============================================================================
# DOCX EXPORT
# =============================================================================

def generate_agenda_docx(agenda: dict) -> bytes:
    """
    Generate a pre-bid meeting agenda as a DOCX document.

    Args:
        agenda: Output of build_meeting_agenda().

    Returns:
        DOCX content as bytes.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Title
    doc.add_heading(f"Pre-bid Meeting Agenda: {agenda.get('project_name', '')}", level=0)
    doc.add_paragraph(f"Generated: {agenda.get('generated_at', '')[:16]}")

    summary = agenda.get("summary", {})
    doc.add_paragraph(
        f"Total items: {summary.get('total_items', 0)} | "
        f"High priority: {summary.get('high_count', 0)} | "
        f"Assigned: {summary.get('assigned_count', 0)}"
    )
    doc.add_paragraph("")

    # Sections
    for section in agenda.get("sections", []):
        doc.add_heading(section["title"], level=1)

        items = section.get("items", [])
        if not items:
            doc.add_paragraph("No items in this section.")
            continue

        # Table
        headers = ["#", "ID", "Title", "Assigned To", "Due Date", "Status"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light List Accent 1'
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for idx, item in enumerate(items):
            row = table.add_row()
            row.cells[0].text = str(idx + 1)
            row.cells[1].text = str(item.get("id", ""))[:20]
            row.cells[2].text = str(item.get("title", ""))[:60]
            row.cells[3].text = str(item.get("assigned_to", ""))[:30]
            row.cells[4].text = str(item.get("due_date", ""))[:12]
            row.cells[5].text = str(item.get("status", ""))[:15]

        doc.add_paragraph("")

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph("Pre-bid Meeting Agenda — Generated by xBOQ Bid Engineer")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# =============================================================================
# PDF EXPORT
# =============================================================================

def generate_agenda_pdf(agenda: dict) -> bytes:
    """
    Generate a pre-bid meeting agenda as a PDF document.

    Args:
        agenda: Output of build_meeting_agenda().

    Returns:
        PDF content as bytes.
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    )
    from reportlab.lib.enums import TA_CENTER

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'AgendaTitle', parent=styles['Heading1'],
        fontSize=18, spaceAfter=12, alignment=TA_CENTER,
        textColor=colors.HexColor('#1a365d'),
    )
    h1_style = ParagraphStyle(
        'AgendaH1', parent=styles['Heading1'],
        fontSize=13, spaceBefore=14, spaceAfter=6,
        textColor=colors.HexColor('#2c5282'),
    )
    normal_style = ParagraphStyle(
        'AgendaNormal', parent=styles['Normal'],
        fontSize=9, spaceAfter=4,
    )
    footer_style = ParagraphStyle(
        'AgendaFooter', parent=styles['Normal'],
        fontSize=8, textColor=colors.HexColor('#718096'), alignment=TA_CENTER,
    )

    table_style = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#edf2f7')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e0')),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
    ])

    def _safe(val, max_len=60):
        s = str(val) if val is not None else ""
        s = s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        return s[:max_len] + "..." if len(s) > max_len else s

    content = []

    # Title
    content.append(Paragraph(
        f"Pre-bid Meeting Agenda: {_safe(agenda.get('project_name', ''), 80)}",
        title_style,
    ))
    content.append(Spacer(1, 6))
    content.append(Paragraph(
        f"Generated: {agenda.get('generated_at', '')[:16]}",
        normal_style,
    ))

    summary = agenda.get("summary", {})
    content.append(Paragraph(
        f"Total items: {summary.get('total_items', 0)} | "
        f"High priority: {summary.get('high_count', 0)} | "
        f"Assigned: {summary.get('assigned_count', 0)}",
        normal_style,
    ))
    content.append(Spacer(1, 12))

    # Sections
    for section in agenda.get("sections", []):
        content.append(Paragraph(section["title"], h1_style))

        items = section.get("items", [])
        if not items:
            content.append(Paragraph("No items in this section.", normal_style))
            continue

        data = [["#", "ID", "Title", "Assigned", "Due", "Status"]]
        for idx, item in enumerate(items):
            data.append([
                Paragraph(str(idx + 1), normal_style),
                Paragraph(_safe(item.get("id", ""), 20), normal_style),
                Paragraph(_safe(item.get("title", ""), 50), normal_style),
                Paragraph(_safe(item.get("assigned_to", ""), 25), normal_style),
                Paragraph(_safe(item.get("due_date", ""), 12), normal_style),
                Paragraph(_safe(item.get("status", ""), 15), normal_style),
            ])

        t = Table(data, colWidths=[25, 60, 170, 80, 60, 60])
        t.setStyle(table_style)
        content.append(t)
        content.append(Spacer(1, 10))

    # Footer
    content.append(Spacer(1, 20))
    content.append(Paragraph(
        "Pre-bid Meeting Agenda — Generated by xBOQ Bid Engineer",
        footer_style,
    ))

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=18 * mm, rightMargin=18 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
    )
    doc.build(content)
    return buffer.getvalue()
